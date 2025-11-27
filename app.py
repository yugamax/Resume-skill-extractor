from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import io
import asyncio
from concurrent.futures import ThreadPoolExecutor
import os
import logging
import ai_parser

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("resume-extractor")

app = FastAPI(title="AI-driven Resume Skill Extractor")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Thread pool for blocking work (model SDK)
executor = ThreadPoolExecutor(max_workers=2)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except Exception as e:
        raise RuntimeError(f"PdfReader import failed: {e}")
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n".join(pages)
    except Exception as e:
        raise RuntimeError(f"PDF parsing failed: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except Exception as e:
        raise RuntimeError(f"docx import failed: {e}")
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"DOCX parsing failed: {e}")

@app.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Accepts .pdf or .docx resume files and returns extracted skills + parse info."""
    filename = file.filename or "unknown"
    ext = filename.split('.')[-1].lower()
    try:
        contents = await file.read()
    except Exception as e:
        logger.exception("Failed to read uploaded file")
        raise HTTPException(status_code=400, detail=f"Failed to read uploaded file: {e}")

    # Extract text depending on extension
    try:
        if ext == 'pdf':
            text = extract_text_from_pdf(contents)
        elif ext == 'docx':
            text = extract_text_from_docx(contents)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .pdf or .docx")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Document parsing error")
        raise HTTPException(status_code=500, detail=str(e))

    # Run the (potentially blocking) AI parser in the threadpool
    loop = asyncio.get_event_loop()
    try:
        parsed = await loop.run_in_executor(executor, ai_parser.extract_resume_info, text)
    except Exception as e:
        logger.exception("AI parsing executor raised an exception")
        return JSONResponse(status_code=500, content={
            "filename": filename,
            "extension": ext,
            "skills": [],
            "error": f"AI parsing execution failed: {e}",
            "ai_raw": None
        })

    # Normalize results
    if not parsed or not isinstance(parsed, dict):
        logger.error("AI parser returned no usable result or invalid type")
        return JSONResponse(status_code=502, content={
            "filename": filename,
            "extension": ext,
            "skills": [],
            "error": "AI parser returned no usable result",
            "ai_raw": None
        })

    # Ensure skills key exists and is a list
    skills = parsed.get("skills")
    if not isinstance(skills, list):
        skills = []

    if parsed.get("error"):
        logger.warning("AI parser reported an error: %s", parsed.get("error"))
        return JSONResponse(status_code=502, content={
            "filename": filename,
            "extension": ext,
            "skills": skills,
            "error": f"AI parsing error: {parsed.get('error')}",
            "ai_raw": parsed.get("raw"),
            "client_used": parsed.get("client_used")
        })

    # Successful response
    is_resume_val = parsed.get("is_resume", False)
    is_resume = True if is_resume_val else False
    response_payload = {
        "filename": filename,
        "extension": ext,
        "skills": skills,
        "is_resume": is_resume
    }
    return JSONResponse(content=response_payload)


@app.get("/health")
async def health():
    """Simple health endpoint to check the service."""
    # check quickly if any Groq client is configured (non-blocking)
    clients = ai_parser.get_available_clients()
    client_count = len(clients) if clients is not None else 0
    note = f"AI-driven skill extraction (configured Groq clients: {client_count})."
    return {"status": "ok", "note": note}


if __name__ == '__main__':
    import uvicorn
    uvicorn.run('app:app', host=os.getenv("HOST", "127.0.0.1"), port=int(os.getenv("PORT", "8000")), reload=True)
